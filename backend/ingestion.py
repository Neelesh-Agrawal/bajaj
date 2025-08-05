"""
Document Ingestion Module - Fixed Version
Handles processing of PDF and DOCX documents from URLs or file uploads
"""

import os
import requests
import tempfile
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional
import asyncio
from urllib.parse import urlparse

# Use basic imports and fallbacks
try:
    from langchain_community.document_loaders import PyPDFLoader
    LANGCHAIN_AVAILABLE = True
except ImportError:
    LANGCHAIN_AVAILABLE = False
    logging.warning("LangChain not available, using basic PDF processing")

try:
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    TEXT_SPLITTER_AVAILABLE = True
except ImportError:
    TEXT_SPLITTER_AVAILABLE = False
    logging.warning("LangChain text splitter not available, using basic splitter")

# Import our utilities
from utils import extract_text_from_pdf, clean_text

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class SimpleTextSplitter:
    """Simple text splitter fallback"""
    
    def __init__(self, chunk_size=1000, chunk_overlap=200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
    
    def split_text(self, text: str) -> List[str]:
        """Split text into chunks"""
        if not text:
            return []
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + self.chunk_size
            chunk = text[start:end]
            
            # Find a good breaking point (sentence end)
            if end < len(text):
                last_period = chunk.rfind('.')
                last_newline = chunk.rfind('\n')
                break_point = max(last_period, last_newline)
                
                if break_point > start + self.chunk_size // 2:
                    chunk = text[start:start + break_point + 1]
                    end = start + break_point + 1
            
            chunks.append(chunk.strip())
            start = end - self.chunk_overlap
            
            if start >= len(text):
                break
        
        return [chunk for chunk in chunks if len(chunk.strip()) > 50]

class DocumentIngestion:
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        """
        Initialize document ingestion with configurable chunking parameters
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
        # Initialize text splitter
        if TEXT_SPLITTER_AVAILABLE:
            self.text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                length_function=len,
                separators=["\n\n", "\n", " ", ""]
            )
        else:
            self.text_splitter = SimpleTextSplitter(chunk_size, chunk_overlap)
        
        # Create temp directory for downloaded files
        self.temp_dir = Path(tempfile.mkdtemp(prefix="doc_ingestion_"))
        logger.info(f"📁 Created temp directory: {self.temp_dir}")
    
    async def download_document(self, url: str) -> str:
        """
        Download document from URL to temporary file
        """
        try:
            logger.info(f"📥 Downloading document from: {url[:50]}...")
            
            # Parse URL to get filename
            parsed_url = urlparse(url)
            filename = Path(parsed_url.path).name
            
            # If no filename in URL, determine from content-type
            if not filename or '.' not in filename:
                try:
                    response = requests.head(url, timeout=30)
                    content_type = response.headers.get('content-type', '').lower()
                    
                    if 'pdf' in content_type:
                        filename = "document.pdf"
                    elif 'word' in content_type or 'docx' in content_type:
                        filename = "document.docx"
                    else:
                        filename = "document.pdf"  # default assumption
                except:
                    filename = "document.pdf"
            
            # Download file
            file_path = self.temp_dir / filename
            
            response = requests.get(url, timeout=60, stream=True)
            response.raise_for_status()
            
            with open(file_path, 'wb') as f:
                for chunk in response.iter_content(chunk_size=8192):
                    if chunk:
                        f.write(chunk)
            
            file_size = file_path.stat().st_size
            logger.info(f"✅ Downloaded: {filename} ({file_size} bytes)")
            
            if file_size == 0:
                raise Exception("Downloaded file is empty")
            
            return str(file_path)
            
        except Exception as e:
            logger.error(f"❌ Error downloading document: {e}")
            raise Exception(f"Failed to download document from {url}: {str(e)}")
    
    def extract_text_from_file(self, file_path: str) -> str:
        """
        Extract text from PDF or DOCX file
        """
        try:
            file_extension = Path(file_path).suffix.lower()
            logger.info(f"📄 Processing {file_extension} file: {Path(file_path).name}")
            
            if file_extension == '.pdf':
                if LANGCHAIN_AVAILABLE:
                    try:
                        loader = PyPDFLoader(file_path)
                        documents = loader.load()
                        text = "\n".join([doc.page_content for doc in documents])
                    except Exception as e:
                        logger.warning(f"LangChain PDF loader failed: {e}, using fallback")
                        text = extract_text_from_pdf(file_path)
                else:
                    text = extract_text_from_pdf(file_path)
                
            elif file_extension in ['.docx', '.doc']:
                try:
                    import docx
                    doc = docx.Document(file_path)
                    text = ""
                    for paragraph in doc.paragraphs:
                        text += paragraph.text + "\n"
                    
                    # Extract text from tables
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                text += cell.text + " "
                            text += "\n"
                            
                except ImportError:
                    raise Exception("python-docx not installed. Install with: pip install python-docx")
                
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
            
            # Clean and normalize text
            text = clean_text(text)
            
            if not text.strip():
                raise Exception("No text content extracted from document")
            
            logger.info(f"✅ Extracted {len(text)} characters from {Path(file_path).name}")
            return text
            
        except Exception as e:
            logger.error(f"❌ Error extracting text from {file_path}: {e}")
            raise Exception(f"Failed to extract text: {str(e)}")
    
    def split_text_into_chunks(self, text: str) -> List[str]:
        """
        Split text into overlapping chunks for processing
        """
        try:
            if not text.strip():
                logger.warning("⚠️ Empty text provided for chunking")
                return []
            
            chunks = self.text_splitter.split_text(text)
            logger.info(f"✂️ Split text into {len(chunks)} chunks")
            
            # Filter out very short chunks
            filtered_chunks = [chunk for chunk in chunks if len(chunk.strip()) > 50]
            
            if not filtered_chunks:
                # If all chunks are too short, create one chunk from the full text
                logger.warning("⚠️ All chunks too short, using full text as single chunk")
                filtered_chunks = [text[:self.chunk_size]]
            
            logger.info(f"📋 Kept {len(filtered_chunks)} chunks after filtering")
            return filtered_chunks
            
        except Exception as e:
            logger.error(f"❌ Error splitting text: {e}")
            # Fallback: create simple chunks
            chunk_size = self.chunk_size
            chunks = [text[i:i+chunk_size] for i in range(0, len(text), chunk_size//2)]
            return [chunk for chunk in chunks if len(chunk.strip()) > 50]
    
    async def process_document(self, document_url: str) -> Dict[str, Any]:
        """
        Process a document from URL: download, extract text, and chunk
        """
        file_path = None
        try:
            logger.info(f"🔄 Starting document processing for: {document_url}")
            
            # Step 1: Download document
            file_path = await self.download_document(document_url)
            
            # Step 2: Extract text
            text = self.extract_text_from_file(file_path)
            
            if not text.strip():
                raise Exception("No text content extracted from document")
            
            # Step 3: Split into chunks
            chunks = self.split_text_into_chunks(text)
            
            if not chunks:
                raise Exception("Failed to create text chunks")
            
            # Step 4: Determine file type
            file_type = Path(file_path).suffix.lower().replace('.', '')
            
            # Step 5: Calculate statistics
            total_chars = len(text)
            avg_chunk_size = sum(len(chunk) for chunk in chunks) / len(chunks) if chunks else 0
            
            result = {
                "source_url": document_url,
                "file_type": file_type,
                "total_characters": total_chars,
                "total_chunks": len(chunks),
                "avg_chunk_size": int(avg_chunk_size),
                "chunks": chunks,
                "processing_status": "success"
            }
            
            logger.info(f"✅ Document processing completed:")
            logger.info(f"   📊 Total characters: {total_chars}")
            logger.info(f"   🧩 Total chunks: {len(chunks)}")
            logger.info(f"   📏 Average chunk size: {int(avg_chunk_size)}")
            
            return result
            
        except Exception as e:
            logger.error(f"❌ Document processing failed: {e}")
            return {
                "source_url": document_url,
                "processing_status": "failed",
                "error": str(e),
                "chunks": [],
                "total_chunks": 0
            }
        finally:
            # Clean up temporary file
            if file_path and os.path.exists(file_path):
                try:
                    os.remove(file_path)
                    logger.info(f"🗑️ Cleaned up temporary file: {Path(file_path).name}")
                except Exception as cleanup_error:
                    logger.warning(f"⚠️ Could not clean up temp file: {cleanup_error}")
    
    def cleanup(self):
        """Clean up temporary directory"""
        try:
            import shutil
            shutil.rmtree(self.temp_dir)
            logger.info(f"🗑️ Cleaned up temp directory: {self.temp_dir}")
        except Exception as e:
            logger.warning(f"⚠️ Could not clean up temp directory: {e}")

# Global instance
document_ingestion = DocumentIngestion()