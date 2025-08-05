"""
Utility Functions for Document Processing
Missing utilities referenced in ingestion.py
"""

import re
import io
import logging
from typing import Optional
from fastapi import UploadFile
import PyPDF2

logger = logging.getLogger(__name__)

def split_text_into_chunks(text: str, chunk_size: int = 500, overlap: int = 50) -> list:
    """
    Split text into overlapping chunks.
    
    Args:
        text: Input text to split
        chunk_size: Size of each chunk (default: 500 characters)
        overlap: Number of overlapping characters between chunks (default: 50)
        
    Returns:
        List of text chunks
    """
    if not text:
        return []

    chunks = []
    start = 0
    text_length = len(text)
    
    while start < text_length:
        end = min(start + chunk_size, text_length)
        chunk = text[start:end].strip()
        chunks.append(chunk)
        start += chunk_size - overlap  # move forward with overlap

    return chunks


def clean_text(text: str) -> str:
    """
    Clean and normalize extracted text
    
    Args:
        text: Raw text to clean
        
    Returns:
        Cleaned text
    """
    if not text:
        return ""
    
    # Remove excessive whitespace
    text = re.sub(r'\s+', ' ', text)
    
    # Remove special characters but keep basic punctuation
    text = re.sub(r'[^\w\s.,!?;:()\-"/]', '', text)
    
    # Remove multiple consecutive punctuation
    text = re.sub(r'[.]{2,}', '.', text)
    text = re.sub(r'[,]{2,}', ',', text)
    
    # Normalize line breaks
    text = text.replace('\n', ' ').replace('\r', ' ')
    
    # Remove extra spaces
    text = ' '.join(text.split())
    
    return text.strip()

def extract_text_from_pdf(file_path_or_upload) -> str:
    """
    Extract text from PDF file
    
    Args:
        file_path_or_upload: File path string or UploadFile object
        
    Returns:
        Extracted text content
    """
    try:
        if isinstance(file_path_or_upload, str):
            # File path
            with open(file_path_or_upload, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        else:
            # UploadFile object
            pdf_content = file_path_or_upload.file.read()
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(pdf_content))
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        
        return clean_text(text)
        
    except Exception as e:
        logger.error(f"Error extracting PDF text: {e}")
        raise Exception(f"Failed to extract PDF text: {str(e)}")

async def extract_text_from_pdf_async(uploaded_file: UploadFile) -> str:
    """
    Async version of PDF text extraction for UploadFile
    
    Args:
        uploaded_file: FastAPI UploadFile object
        
    Returns:
        Extracted text content
    """
    try:
        # Read file content
        content = await uploaded_file.read()
        
        # Reset file pointer
        await uploaded_file.seek(0)
        
        # Extract text using PyPDF2
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
        text = ""
        
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        
        return clean_text(text)
        
    except Exception as e:
        logger.error(f"Error extracting PDF text async: {e}")
        raise Exception(f"Failed to extract PDF text: {str(e)}")

def extract_text_from_docx_bytes(docx_bytes: bytes) -> str:
    """
    Extract text from DOCX bytes
    
    Args:
        docx_bytes: DOCX file content as bytes
        
    Returns:
        Extracted text content
    """
    try:
        import docx
        from io import BytesIO
        
        # Create a BytesIO object from bytes
        docx_file = BytesIO(docx_bytes)
        
        # Load document
        doc = docx.Document(docx_file)
        
        # Extract text from paragraphs
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        return clean_text(text)
        
    except ImportError:
        logger.error("python-docx not installed. Install with: pip install python-docx")
        raise Exception("DOCX processing not available. Install python-docx package.")
    except Exception as e:
        logger.error(f"Error extracting DOCX text: {e}")
        raise Exception(f"Failed to extract DOCX text: {str(e)}")

def validate_file_type(filename: str) -> bool:
    """
    Validate if file type is supported
    
    Args:
        filename: Name of the file
        
    Returns:
        True if supported, False otherwise
    """
    supported_extensions = ['.pdf', '.docx', '.doc']
    file_extension = filename.lower().split('.')[-1]
    return f'.{file_extension}' in supported_extensions

def get_file_size_mb(file_size_bytes: int) -> float:
    """
    Convert file size from bytes to MB
    
    Args:
        file_size_bytes: File size in bytes
        
    Returns:
        File size in MB
    """
    return file_size_bytes / (1024 * 1024)

def truncate_text(text: str, max_length: int = 1000) -> str:
    """
    Truncate text to specified length with ellipsis
    
    Args:
        text: Text to truncate
        max_length: Maximum length
        
    Returns:
        Truncated text
    """
    if len(text) <= max_length:
        return text
    return text[:max_length-3] + "..."

def extract_key_phrases(text: str) -> list:
    """
    Extract key phrases from text (simple implementation)
    
    Args:
        text: Input text
        
    Returns:
        List of key phrases/words
    """
    # Simple keyword extraction
    insurance_keywords = [
        'premium', 'coverage', 'deductible', 'claim', 'policy', 
        'benefit', 'exclusion', 'waiting period', 'co-pay',
        'pre-existing', 'maternity', 'hospitalization'
    ]
    
    text_lower = text.lower()
    found_keywords = []
    
    for keyword in insurance_keywords:
        if keyword in text_lower:
            found_keywords.append(keyword)
    
    return found_keywords

# Add any other utility functions that might be referenced
def format_confidence_score(confidence: float) -> str:
    """Format confidence score as percentage"""
    return f"{confidence * 100:.1f}%"

def sanitize_filename(filename: str) -> str:
    """Sanitize filename for safe storage"""
    # Remove or replace unsafe characters
    sanitized = re.sub(r'[^\w\-_\.]', '_', filename)
    return sanitized[:100]  # Limit length